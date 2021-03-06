from datetime import date
from io import BytesIO
from typing import Iterable

from src.enums import TransactionStatus
from src.repositories import TransactionRepository

from src.models import DailyReport, TransactionCreate, MarketTransactionCreate, ResourceTransactionCreate, ServiceTransactionCreate
from src.schemas import Transaction, MarketTransaction, ServiceTransaction, ResourceTransaction, Transactions_Resources_Services
import src.models as models

import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

class TransactionService:
    def __init__(self, transactionRepository: TransactionRepository):
        self.transactionRepository = transactionRepository
    
    def findAll(self) -> Iterable[models.Transaction]:
        return self.transactionRepository.findAll()

    def findAllMarketTransactions(self) -> Iterable[models.MarketTransaction]:
        return self.transactionRepository.findAllMarketTransactions()

    def findAllServiceTransactions(self) -> Iterable[models.ServiceTransaction]:
        return self.transactionRepository.findAllServiceTransactions()

    def findAllResourceTransactions(self) -> Iterable[models.ResourceTransaction]:
        return self.transactionRepository.findAllResourceTransactions()

    def createInvoice(self, transactionId: int):
        tx = self.findById(transactionId)
        stream = BytesIO()
        document = docx.Document()
        
        document.add_heading('Transaction ' + str(tx.id) + ': ' + tx.type.capitalize(), 0)
        paragraph = document.add_paragraph('User: ' + str(tx.userId) + "")
        paragraph = document.add_paragraph('Status: ' + str(tx.status))
        paragraph = document.add_paragraph('Total price: ' + str(tx.price) + '$')
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        document.save(stream)
        stream.flush()
        stream.seek(0)
        return stream
    
    def findById(self, id: int) -> models.Transaction:
        return self.transactionRepository.findById(id)

    def findAllByCreatedAt(self, createdAt: date) -> Iterable[models.Transaction]:
        return self.transactionRepository.findAllByCreatedAt(createdAt)

    def findAllByUserId(self, userId: int) -> Iterable[models.Transaction]:
        return self.transactionRepository.findByUserId(userId)

    def pay(self, tx: any):
        if type(tx) is MarketTransactionCreate:
            self.createMarket(tx)
        elif type(tx) is ServiceTransactionCreate:
            self.createService(tx)
        elif type(tx) is ResourceTransactionCreate:
            self.createResource(tx)
        return 200

    def createBase(self, tx: TransactionCreate) -> models.Transaction:
        return self.transactionRepository.create(
            Transaction(
                userId = tx.userId, 
                type = "base", 
                createdAt = date.today(), 
                price = tx.price, 
                status = TransactionStatus.PENDING,
            )
        )

    def createMarket(self, tx: MarketTransactionCreate) -> models.MarketTransaction:
        return self.transactionRepository.create(
            MarketTransaction(
                userId = tx.userId, 
                type = "market", 
                createdAt = date.today(), 
                price = tx.price, 
                status = TransactionStatus.PENDING,
                shoppingCartId = tx.shoppingCartId
            )
        )

    def createService(self, tx: ServiceTransactionCreate) -> models.ServiceTransaction:
        return self.transactionRepository.create(
            ServiceTransaction(
                userId = tx.userId, 
                type = "service", 
                createdAt = date.today(), 
                price = tx.price, 
                status = TransactionStatus.PENDING,
                # petId = tx.petId,
                serviceIds = [Transactions_Resources_Services(data_id = serviceId) for serviceId in tx.serviceIds]
            )
        )

    def createResource(self, tx: ResourceTransactionCreate) -> models.ResourceTransaction:
        return self.transactionRepository.create(
            ResourceTransaction(
                userId = tx.userId, 
                type = "resource", 
                createdAt = date.today(), 
                price = tx.price, 
                status = TransactionStatus.PENDING,
                # petId = tx.petId,
                resourceIds = [Transactions_Resources_Services(data_id = resourceId) for resourceId in tx.resourceIds]
            )
        )

    def setTransactionStatus(self, id: int, status: TransactionStatus) -> bool:
        transaction = self.findById(id)
        transaction.status = status
        self.transactionRepository.update(transaction)
        return True

    def cancelTransaction(self, id: int) -> bool:
        self.setStatus(id, TransactionStatus.CANCELED)
        return True

    def getDailyReportForDate(self, date: date) -> DailyReport:
        txs = self.findAllByCreatedAt(date)
        return DailyReport(
            date = date, 
            totalTransactions = len(txs), 
            totalPending = len(list(filter(lambda x: x.status == TransactionStatus.PENDING, txs))), 
            totalResolved = len(list(filter(lambda x: x.status == TransactionStatus.RESOLVED, txs))), 
            totalCanceled = len(list(filter(lambda x: x.status == TransactionStatus.CANCELED, txs))), 
            totalPrice = sum(map(lambda x: x.price, txs))
        )